#!/usr/bin/env python
# coding: utf-8

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from random import randrange
import os


# 加载配置文件信息
def load_conf(config_name):
    config_dic = {}
    with open(config_name, 'r') as f:
        lines = f.readlines()
        for config in lines:
            config_str = config.strip('\n')
            config_key = config_str.split('=')[0]
            config_value = config_str.split('=')[1]
            config_dic[config_key] = config_value
    return config_dic


# 生成减法列表
def get_sub_list():
    sub_li = []
    for x in range(int(configs_dic['range'])):
        for y in range(x + 1):
            sub_item = str(x) + " - " + str(y) + " = "
            sub_li.append(sub_item)
    return sub_li


# 生成加法列表
def get_add_list():
    add_li = []
    for a in range(int(configs_dic['range'])):
        for b in range(int(configs_dic['range'])):
            add_item = str(a) + " + " + str(b) + " = "
            add_li.append(add_item)
    return add_li


# 生成目标列表
def get_target_list(add_li, sub_li):
    all_li = add_li + sub_li
    target_list = []
    for i in range(int(configs_dic['page'])):
        random_index = randrange(0, len(all_li))
        target_list.append(all_li[random_index])
    return target_list


# 把目标列表写入word
def write_to_word(target_list):
    row_size = int(len(target_list) // 3)
    left_target_list = target_list[:row_size]
    middle_target_list = target_list[row_size:row_size * 2]
    right_target_list = target_list[row_size * 2:]
    # 打开文档
    document = Document(docx=os.path.join(os.getcwd(), 'default.docx'))

    # 加入标题
    document.add_heading(u'10以内加减,加强练习', 0)

    # 添加文本

    for i in range(row_size):
        paragraph = document.add_paragraph()
        # 设置中文字体
        run = paragraph.add_run(
            left_target_list[i] + "    " + middle_target_list[i] + "     " + right_target_list[i] + "     ")
        run.font.name = u'宋体'
        run.font.size = Pt(20)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 保存文件
    document.save(u'加减.docx')


if __name__ == '__main__':
    config_file_name = './conf/conf.ini'
    configs_dic = load_conf(config_file_name)
    add_list = get_add_list()
    sub_list = get_sub_list()
    target_list_all = get_target_list(add_list, sub_list)
    write_to_word(target_list_all)
